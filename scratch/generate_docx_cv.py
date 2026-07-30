import os
import subprocess
import sys

# Self-heal dependency: install python-docx if not present
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cv():
    doc = docx.Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Color Palette (Slate Blue / Dark Gray)
    PRIMARY_COLOR = RGBColor(15, 23, 42)    # Slate Blue #0f172a
    SECONDARY_COLOR = RGBColor(59, 130, 246) # Blue #3b82f6
    TEXT_COLOR = RGBColor(51, 65, 85)       # Slate Gray #334155
    
    # Styles config
    # Normal Text Style
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(10.5)
    font_normal.color.rgb = TEXT_COLOR
    
    # Helper to add section headers
    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        
        # Add a subtle horizontal line border below heading
        pBorder = doc.add_paragraph()
        pBorder.paragraph_format.space_before = Pt(0)
        pBorder.paragraph_format.space_after = Pt(8)
        runBorder = pBorder.add_run("―" * 70)
        runBorder.font.size = Pt(6)
        runBorder.font.color.rgb = RGBColor(226, 232, 240)

    # Helper to add job titles
    def add_job_header(company, title, dates, location):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        
        run_title = p.add_run(f"{title}  |  ")
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = PRIMARY_COLOR
        
        run_comp = p.add_run(f"{company}")
        run_comp.bold = True
        run_comp.font.size = Pt(11)
        run_comp.font.color.rgb = SECONDARY_COLOR
        
        # Add float right dates and location
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(4)
        p_sub.paragraph_format.keep_with_next = True
        
        run_details = p_sub.add_run(f"📅 {dates}   •   📍 {location}")
        run_details.font.size = Pt(9.5)
        run_details.italic = True
        run_details.font.color.rgb = RGBColor(100, 116, 139)

    # 1. Header (Name, Title, Contacts)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("ANTONIO GUTIÉRREZ JIMÉNEZ")
    run_name.font.size = Pt(20)
    run_name.bold = True
    run_name.font.color.rgb = PRIMARY_COLOR
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_after = Pt(6)
    run_sub = p_subtitle.add_run("B2B Sales Executive  |  Fintech & Payments Senior Specialist")
    run_sub.font.size = Pt(11)
    run_sub.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)
    run_contact = p_contact.add_run("📍 Cancún, Quintana Roo, México   |   📞 +52 998 119 1903\n✉️ antoniogtzjimenez@gmail.com   |   🔗 linkedin.com/in/agjbusiness/")
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = RGBColor(71, 85, 105)

    # 2. Resumen Profesional
    add_section_header("Resumen Profesional")
    p_resumen = doc.add_paragraph(
        "Ejecutivo de Ventas B2B Senior y especialista en Adquirencia (Merchant Acquiring) y Tecnología de Pagos (PayTech), "
        "con más de 5 años de trayectoria acelerando el crecimiento comercial de grandes corporaciones fintech y multinacionales en México y LATAM. "
        "Reconocido por un enfoque consultivo de alta eficiencia: priorizo la adquisición de cuentas estratégicas (Middle Market/High Potential) "
        "de alto volumen transaccional, optimizando el costo de onboarding e integraciones tecnológicas (APIs/ISVs). Experto en prospección "
        "outbound autónoma en campo y corporativa, con un portafolio activo de más de 3,000 conexiones en la industria de pagos digitales en LATAM."
    )
    p_resumen.paragraph_format.space_after = Pt(10)

    # 3. Areas de Expertise
    add_section_header("Áreas de Expertise")
    skills = [
        ("Merchant Acquiring & PayTech", "Procesamiento de pagos, adquirencia, pasarelas de pago (gateways), prevención de fraude, transacciones transfronterizas."),
        ("Ventas Consultivas & Outbound Hunting", "Prospección activa, mapeo de territorios, ciclos de venta complejos B2B de ciclo largo."),
        ("Integraciones API & ISV Partnerships", "Alianzas comerciales y técnicas con ERPs, PMS y sistemas de punto de venta (POS) para automatización de comercios."),
        ("Sales Ops & Inteligencia Comercial", "Creación de herramientas propias de SalesTech, automatizaciones con Power BI / Power Automate, Salesforce y CRM.")
    ]
    for skill_title, skill_desc in skills:
        p_skill = doc.add_paragraph(style='List Bullet')
        p_skill.paragraph_format.space_after = Pt(3)
        run_st = p_skill.add_run(f"{skill_title}: ")
        run_st.bold = True
        p_skill.add_run(skill_desc)

    # 4. Experiencia Profesional
    add_section_header("Experiencia Profesional")
    
    # -- LATAM Payments & eCommerce
    add_job_header("LATAM Payments & eCommerce", "Co-Founder", "05/2024 – Presente", "LATAM · Remoto")
    p1 = doc.add_paragraph("Cofundé y coordino una comunidad activa de más de 500 profesionales de pagos digitales, adquirencia y comercio electrónico en México y LATAM.", style='List Bullet')
    p1.paragraph_format.space_after = Pt(3)
    p2 = doc.add_paragraph("Lidero la creación de contenido y curaduría de mejores prácticas del sector, conectando ISOs, PSPs, fintechs y ejecutivos comerciales para acelerar oportunidades de negocio regionales (ABM).", style='List Bullet')
    p2.paragraph_format.space_after = Pt(8)

    # -- Fiserv
    add_job_header("Fiserv", "Business Advisor", "02/2025 – 10/2025", "Cancún, México")
    p3 = doc.add_paragraph("Gestioné la retención, mitigación de Churn y desarrollo de una cartera activa de más de 80 comercios corporativos (merchants) de gran escala.", style='List Bullet')
    p3.paragraph_format.space_after = Pt(3)
    p4 = doc.add_paragraph("Diseñé e implementé un modelo dinámico de monitoreo transaccional Month-over-Month (MoM) mediante Power BI, optimizando la toma de decisiones comerciales.", style='List Bullet')
    p4.paragraph_format.space_after = Pt(3)
    p5 = doc.add_paragraph("Desarrollé un workflow automatizado de reactivación mediante Power Automate e integraciones de Salesforce, generando un promedio de +15 oportunidades de negocio mensuales desde la alianza bancaria asignada.", style='List Bullet')
    p5.paragraph_format.space_after = Pt(8)

    # -- Clip
    add_job_header("Clip", "Asesor Comercial - Middle Market / High Potential", "07/2021 – 02/2025", "Cancún, México")
    bullets_clip = [
        "Top Performer & Cumplimiento: Posicionado en el Top 12% nacional (Lugar #22 de 184 ejecutivos de Middle Market) en H1 2022, superando las cuotas de volumen mensual asignadas por más del 280% de forma consistente ($2.8M a $5.8M MXN promedio frente a la meta de $1M) y logrando el 3er lugar general del podio nacional.",
        "Eficiencia de Cartera (High Value): Diseñé y ejecuté una estrategia comercial enfocada en cuentas medianas de alto potencial, logrando un TPV promedio por deal de $555k MXN (60% superior a la media del segmento), maximizando el volumen procesado con una fracción del costo operativo de integración y soporte.",
        "Cierre de Cuentas Enterprise (Outbound): Cerré de manera autónoma las cuentas de mayor volumen de la cartera en el sector turismo de lujo, destacando un operador de yates de lujo ($14.5M MXN YTD TPV) y una firma de turismo de aventura ($20.0M MXN YTD TPV) mediante prospección activa en frío y cambaceo estratégico.",
        "Integraciones Tecnológicas & APIs: Lideré negociaciones comerciales complejas e integraciones de pasarela de pagos vía API/ISV con sistemas clave (Bistrosoft, Profitroom, Odoo ERP), incrementando la retención de clientes a largo plazo con una tasa de churn cercana a cero.",
        "Resultado consolidado de la cartera: $69M MXN de TPV total acumulado, con un 75.3% del volumen total auto-generado vía Outbound."
    ]
    for b in bullets_clip:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.space_after = Pt(3)
        # Bold lead-in
        parts = b.split(":", 1)
        if len(parts) == 2:
            lead = p_b.add_run(parts[0] + ":")
            lead.bold = True
            p_b.add_run(parts[1])
        else:
            p_b.add_run(b)
            
    p_b.paragraph_format.space_after = Pt(8)

    # -- Japan Tobacco International (JTI)
    add_job_header("Japan Tobacco International (JTI)", "Account Executive · Southeast & Bajío", "07/2018 – 12/2020", "Cancún & Aguascalientes, México")
    p6 = doc.add_paragraph("Lidero la expansión territorial y negociación de Key Accounts (KAM) en el sector HORECA, logrando un crecimiento del +40% en Share of Opportunity en Cancún y Riviera Maya vs. el año previo.", style='List Bullet')
    p6.paragraph_format.space_after = Pt(3)
    p7 = doc.add_paragraph("Expandí la cartera de clientes activos en un +35%, asegurando contratos de distribución directa con más de 100 hoteles premium y cadenas internacionales de hospitalidad líderes en la región.", style='List Bullet')
    p7.paragraph_format.space_after = Pt(3)
    p8 = doc.add_paragraph("Coordiné y lideré un equipo comercial de 3 personas en campo (FSF) para la región del Bajío.", style='List Bullet')
    p8.paragraph_format.space_after = Pt(8)


    # 5. Educacion y Certificaciones
    add_section_header("Educación Continua & Certificaciones")
    certs = [
        ("McKinsey Forward Program", "McKinsey.org", "120 Horas", "Especialización en liderazgo adaptativo, resolución estructurada de problemas complejos y metodologías ágiles."),
        ("Growth 101", "Kurios", "30 Horas", "Metodología de crecimiento acelerado y marcos de experimentación ágil de producto."),
        ("Mastering Ventas", "Sales Professional", "70 Horas", "Estructura de equipos comerciales, playbook de ventas B2B y automatización del stack SalesTech."),
        ("Curso SDR Primera Reunión", "LATAM SDR Leaders", "16 Horas", "Estrategias avanzadas de prospección en frío y agendamiento de cuentas Enterprise.")
    ]
    for title, inst, duration, desc in certs:
        p_c = doc.add_paragraph(style='List Bullet')
        p_c.paragraph_format.space_after = Pt(3)
        run_ct = p_c.add_run(f"{title} ")
        run_ct.bold = True
        run_ci = p_c.add_run(f"({inst} • {duration}): ")
        run_ci.italic = True
        p_c.add_run(desc)

    # Output path
    downloads_path = r"C:\Users\Antonio\OneDrive\Downloads"
    if not os.path.exists(downloads_path):
        downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads')
    
    os.makedirs(downloads_path, exist_ok=True)
    output_file = os.path.join(downloads_path, "CV_Antonio_Gutierrez_Jimenez_Confidencial.docx")
    
    doc.save(output_file)

    print(f"CV created successfully at: {output_file}")
    return output_file


if __name__ == "__main__":
    create_cv()
